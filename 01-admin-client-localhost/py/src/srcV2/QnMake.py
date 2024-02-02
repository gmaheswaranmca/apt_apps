import docx
import pandas as pd
import os
import io
import os.path

import random
import string

import mysql.connector

from Captions import Captions
from gfQnMake import gfQnMake

class QnMake:
    @classmethod
    def QnMake(cls):
        caption_srcPath = Captions.QnMake_QnMake_caption_srcPath#!
        caption_srcFileNumber = Captions.QnMake_QnMake_caption_srcFileNumber#!
        caption_srcVersionNumber = Captions.QnMake_QnMake_caption_srcVersionNumber#!
        default_srcPath = 'F:/v2m/techLang/assignment/2020-07-08-Apt-BulkQuestions/automate/qn'
        srcPath = input(caption_srcPath)
        if(srcPath==''):
            srcPath = default_srcPath

        lfile_id = 1
        lfile_version_number = 1

        srcFileNumber = input(caption_srcFileNumber) #'501'
        srcVersionNumber = input(caption_srcVersionNumber) #'1'

        FilePathPrefix = srcPath + '/' + srcFileNumber + '-v' + srcVersionNumber
        qnFilePath = FilePathPrefix + '_qn.docx'
        optFilePath = FilePathPrefix + '_opt.docx'
        outFilePath = FilePathPrefix + '_out.json'
        outFilePathV2 = FilePathPrefix + '_outV2.json'

        quiz_name = input('Test Paper Name:')#'Test 01';
        question_count=input('Question Count:')#40;
        time_limit=input('Time Limit(mins):')#50;

        caption_databaseName = Captions.gfQnMake_getConV1_caption_databaseName#!
        default_databaseName = 't20apt00'
        databaseName = input(caption_databaseName)
        if(databaseName==''):
            databaseName = default_databaseName
        QnMake.QnMakeCaller(qnFilePath, optFilePath, outFilePath, outFilePathV2, quiz_name, question_count, time_limit, databaseName)
    @classmethod
    def QnMakeCaller(cls, qnFilePath, optFilePath, outFilePath, outFilePathV2, quiz_name, question_count, time_limit, databaseName):
        gfQnMake.databaseName = databaseName
        '''caption_srcPath = Captions.QnMake_QnMake_caption_srcPath#!
        caption_srcFileNumber = Captions.QnMake_QnMake_caption_srcFileNumber#!
        caption_srcVersionNumber = Captions.QnMake_QnMake_caption_srcVersionNumber#!
        default_srcPath = 'F:/v2m/techLang/assignment/2020-07-08-Apt-BulkQuestions/automate/qn'
        srcPath = input(caption_srcPath)
        if(srcPath==''):
            srcPath = default_srcPath

        lfile_id = 1
        lfile_version_number = 1

        srcFileNumber = input(caption_srcFileNumber) #'501'
        srcVersionNumber = input(caption_srcVersionNumber) #'1'

        FilePathPrefix = srcPath + '/' + srcFileNumber + '-v' + srcVersionNumber
        qnFilePath = FilePathPrefix + '_qn.docx'
        optFilePath = FilePathPrefix + '_opt.docx'
        outFilePath = FilePathPrefix + '_out.json'
        outFilePathV2 = FilePathPrefix + '_outV2.json' '''

        msg = ''
        msg += qnFilePath + '\n'
        msg += optFilePath + '\n'

        docQn = docx.Document(qnFilePath)
        docOpt = docx.Document(optFilePath)

        docTblRowQn = docQn.tables[0].rows
        docTblRowOpt = docOpt.tables[0].rows

        pmsg = [msg]
        gfQnMake.ValidateDocx(docTblRowQn,pmsg)
        gfQnMake.ValidateDocx(docTblRowOpt,pmsg)
        msg = pmsg[0]
        print(msg)

        canProceed = input('\nAre you sure to continue (y/n)?')
        if canProceed=='n':
            return

        DatQn = []
        DatOpt = []
        DatGrp = ['[]']
        QnSec = []
        DatQnOpt = []	
        DatQnOptV2 = []	

        gfQnMake.ParseQn(docTblRowQn,DatQn,DatGrp,QnSec)
        gfQnMake.ParseOpt(docTblRowOpt,DatOpt)
        gfQnMake.DoQnOpt(DatQn,DatGrp,DatOpt,QnSec,DatQnOpt,DatQnOptV2)

    #	print()
    #	print(DatGrp[0])#,QnSec)

        #print()
        print('v1 Qns:',len(DatQnOpt))
        print('v2 Qns:',len(DatQnOptV2))
        print(gfQnMake.ParseSecShuffle(QnSec))
        
        canProceed = input('\nAre you sure to Save v1 (y/n)?')
        #canProceed = 'y'
        if canProceed=='y':	
            df = pd.DataFrame.from_dict(DatQnOpt)
            df.to_json(outFilePath,orient='records')
            print(outFilePath,'Created')

        #canProceed = input('\nAre you sure to Save v2 (y/n)?')
        canProceed = 'y'
        if canProceed=='y':	
            #df = pd.DataFrame.from_dict(DatQnOptV2)
            #df.to_json(outFilePathV2,orient='records')
            #print(outFilePathV2,'Created')
            import json
            out_file = open(outFilePathV2, "w") 
            json.dump({'question_base':DatQnOptV2,'question_section':QnSec,'question_group':DatGrp[0]},out_file,indent=4)
            
        #canProceed = input('\nBuild DB v1 (y/n)?')
        canProceed = 'y'
        if canProceed=='y':	
            #quiz_name = input('Test Paper Name:')#'Test 01';
            #question_count=input('Question Count:')#40;
            #time_limit=input('Time Limit(mins):')#50;

            gfQnMake.writeQuizV1(DatQnOpt,quiz_name,question_count,time_limit)
    @classmethod
    def InvQnMake(cls):
        QnMake.QnMake()
        canProceed = input('\nAnother Qn File (y/n)?')
        while canProceed=='y':	
            QnMake.QnMake()
            canProceed = input('\nAnother Qn File (y/n)?')